@router.post("/pdf-a-word-ocr-upload/")
def pdf_a_word_ocr_upload(files: list[UploadFile] = File(...)):
    """Convierte archivos PDF escaneados a Word usando OCR"""
    temp_dir = tempfile.mkdtemp()
    archivos_convertidos = []
    
    for file in files:
        try:
            # Verificar que sea un archivo PDF
            if not file.filename.lower().endswith('.pdf'):
                continue
                
            # Guardar archivo PDF temporalmente
            temp_input_path = os.path.join(temp_dir, file.filename)
            with open(temp_input_path, "wb") as f:
                shutil.copyfileobj(file.file, f)
            
            # Crear nombre del archivo de salida
            nombre_archivo_base, _ = os.path.splitext(os.path.basename(file.filename))
            nombre_archivo_docx = nombre_archivo_base + "_ocr.docx"
            archivo_salida = os.path.join(temp_dir, nombre_archivo_docx)
            
            texto_extraido = ""
            
            # Primero intentar extracción normal de texto
            try:
                import pdfplumber
                with pdfplumber.open(temp_input_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            texto_extraido += page_text + "\n\n"
                print(f"✅ Texto extraído normalmente: {len(texto_extraido)} caracteres")
                
            except Exception as e:
                print(f"❌ Error extracción normal: {e}")
            
            # Si no se extrajo suficiente texto, usar OCR
            if len(texto_extraido.strip()) < 100:  # Menos de 100 caracteres = probablemente escaneado
                print("🔍 PDF parece escaneado, aplicando OCR...")
                
                try:
                    # Intentar usar OCR si las dependencias están disponibles
                    try:
                        import pdf2image
                        import pytesseract
                        
                        # Convertir PDF a imágenes
                        images = pdf2image.convert_from_path(temp_input_path)
                        print(f"📄 PDF convertido a {len(images)} imágenes para OCR")
                        
                        texto_ocr = ""
                        for i, image in enumerate(images):
                            print(f"   🔍 Procesando página {i+1} con OCR...")
                            
                            # Extraer texto usando Tesseract OCR
                            page_text = pytesseract.image_to_string(image, lang='spa+eng')
                            
                            if page_text.strip():
                                texto_ocr += f"--- PÁGINA {i+1} ---\n{page_text}\n\n"
                                print(f"   ✅ Página {i+1}: {len(page_text)} caracteres extraídos con OCR")
                            else:
                                print(f"   ⚠️  Página {i+1}: No se pudo extraer texto con OCR")
                        
                        if texto_ocr.strip():
                            texto_extraido = texto_ocr
                            print(f"✅ Texto extraído con OCR: {len(texto_ocr)} caracteres")
                        else:
                            print("❌ OCR no pudo extraer texto")
                            
                    except ImportError:
                        print("❌ Dependencias de OCR no disponibles")
                        texto_extraido = "OCR no disponible. Se requieren: pdf2image, pytesseract, tesseract-ocr"
                        
                except Exception as e:
                    print(f"❌ Error en OCR: {e}")
                    texto_extraido = f"Error en OCR: {str(e)}"
            
            # Crear documento Word con el texto extraído
            try:
                from docx import Document
                doc = Document()
                
                # Agregar título
                doc.add_heading(f'Documento convertido: {file.filename}', 0)
                doc.add_paragraph(f'Archivo original: {file.filename}')
                doc.add_paragraph(f'Fecha de conversión: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}')
                doc.add_paragraph('─' * 50)
                
                if texto_extraido.strip():
                    # Procesar el texto extraído
                    lineas = texto_extraido.split('\n')
                    for linea in lineas:
                        if linea.strip() and not linea.startswith('--- PÁGINA'):
                            doc.add_paragraph(linea.strip())
                else:
                    doc.add_paragraph('No se pudo extraer texto del PDF.')
                    doc.add_paragraph('Posibles causas:')
                    doc.add_paragraph('• El PDF está protegido con contraseña')
                    doc.add_paragraph('• El PDF contiene solo imágenes (escaneado)')
                    doc.add_paragraph('• El archivo está corrupto')
                    doc.add_paragraph('• Formato no compatible')
                
                # Guardar el documento
                doc.save(archivo_salida)
                archivos_convertidos.append(archivo_salida)
                print(f"✅ Documento Word creado: {archivo_salida}")
                
            except ImportError:
                # Si no está disponible python-docx, crear archivo de texto
                archivo_salida_txt = archivo_salida.replace('.docx', '.txt')
                with open(archivo_salida_txt, 'w', encoding='utf-8') as txt_file:
                    txt_file.write(f"DOCUMENTO CONVERTIDO: {file.filename}\n")
                    txt_file.write("=" * 50 + "\n\n")
                    if texto_extraido:
                        txt_file.write(texto_extraido)
                    else:
                        txt_file.write("No se pudo extraer texto del PDF.\n")
                archivos_convertidos.append(archivo_salida_txt)
                print(f"✅ Archivo de texto creado: {archivo_salida_txt}")
                
        except Exception as e:
            print(f"❌ Error general procesando {file.filename}: {str(e)}")
            continue
    
    if not archivos_convertidos:
        return JSONResponse(
            status_code=400,
            content={"error": "No se pudo convertir ningún archivo PDF. Verifique que los archivos sean PDFs válidos."}
        )
    
    # Crear ZIP con los archivos convertidos
    zip_path = os.path.join(temp_dir, "pdf_a_word_ocr_convertidos.zip")
    with zipfile.ZipFile(zip_path, "w") as zip极:
        for archivo in archivos_convertidos:
            zipf.write(archivo, os.path.basename(archivo))
    
    return FileResponse(
        zip_path,
        filename="pdf_a_word_ocr_convertidos.zip",
        media_type="application/zip"
    )
