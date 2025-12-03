@router.post("/pdf-a-word-upload/")
def pdf_a_word_upload(files: list[UploadFile] = File(...)):
    """Convierte archivos PDF a Word (.docx) extrayendo el texto real, incluyendo OCR para PDFs escaneados"""
    temp_dir = tempfile.mkdtemp()
    archivos_convertidos = []
    
    for file in files:
        try:
            # Verificar que sea un archivo PDF
            if not file.filename.lower().endswith('.pdf'):
                continue
                
            # Guardar archivo PDF temporalmente
            temp_input_path = os.path.join(temp_dir, file.filename)
            with open(temp_input_path, "wb") as f:
                shutil.copyfileobj(file.file, f)
            
            # Crear nombre del archivo de salida
            nombre_archivo_base, _ = os.path.splitext(os.path.basename(file.filename))
            nombre_archivo_docx = nombre_archivo_base + ".docx"
            archivo_salida = os.path.join(temp_dir, nombre_archivo_docx)
            
            # Intentar extraer texto del PDF usando múltiples métodos
            texto_extraido = ""
            metodo_usado = ""
            
            try:
                # Método 1: Intentar con pdfplumber (más robusto)
                import pdfplumber
                with pdfplumber.open(temp_input_path) as pdf:
                    total_paginas = len(pdf.pages)
                    print(f"📄 PDF tiene {total_paginas} páginas")
                    
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            texto_extraido += page_text + "\n\n"
                            print(f"✅ Página {i+1}: {len(page_text)} caracteres extraídos")
                        else:
                            print(f"⚠️ Página {i+1}: No se pudo extraer texto (probablemente escaneada)")
                
                if texto_extraido.strip():
                    metodo_usado = "pdfplumber"
                    print(f"✅ Texto extraído con pdfplumber: {len(texto_extraido)} caracteres")
                else:
                    print("⚠️ No se extrajo texto con pdfplumber, intentando OCR...")
                    
            except Exception as e:
                print(f"❌ Error con pdfplumber: {e}")
            
            # Si no se extrajo texto, intentar con PyPDF2
            if not texto_extraido.strip():
                try:
                    import PyPDF2
                    with open(temp_input_path, 'rb') as pdf_file:
                        pdf_reader = PyPDF2.PdfReader(pdf_file)
                        for page_num in range(len(pdf_reader.pages)):
                            page = pdf_reader.pages[page_num]
                            page_text = page.extract_text()
                            if page_text:
                                texto_extraido += page_text + "\n\n"
                    
                    if texto_extraido.strip():
                        metodo_usado = "PyPDF2"
                        print(f"✅ Texto extraído con PyPDF2: {len(texto_extraido)} caracteres")
                    
                except Exception as e2:
                    print(f"❌ Error con PyPDF2: {e2}")
            
            # Si aún no se extrajo texto, intentar OCR
            if not texto_extraido.strip():
                try:
                    print("🔍 Intentando OCR con Tesseract...")
                    import pytesseract
                    from PIL import Image
                    import pdf2image
                    
                    # Convertir PDF a imágenes
                    images = pdf2image.convert_from_path(temp_input_path)
                    print(f"📷 PDF convertido a {len(images)} imágenes")
                    
                    for i, image in enumerate(images):
                        print(f"🔍 Procesando imagen {i+1}/{len(images)} con OCR...")
                        try:
                            # Configurar Tesseract para español
                            ocr_text = pytesseract.image_to_string(image, lang='spa+eng')
                            if ocr_text.strip():
                                texto_extraido += f"--- PÁGINA {i+1} ---\n{ocr_text}\n\n"
                                print(f"✅ Página {i+1}: {len(ocr_text)} caracteres extraídos con OCR")
                            else:
                                print(f"⚠️ Página {i+1}: No se pudo extraer texto con OCR")
                        except Exception as ocr_error:
                            print(f"❌ Error OCR en página {i+1}: {ocr_error}")
                    
                    if texto_extraido.strip():
                        metodo_usado = "OCR (Tesseract)"
                        print(f"✅ Texto extraído con OCR: {len(texto_extraido)} caracteres")
                    
                except ImportError as import_error:
                    print(f"❌ Dependencias OCR no disponibles: {import_error}")
                    print("💡 Instale: pip install pytesseract pdf2image")
                except Exception as ocr_error:
                    print(f"❌ Error general con OCR: {ocr_error}")
            
            # Si aún no se extrajo texto, intentar con LibreOffice
            if not texto_extraido.strip():
                try:
                    print("🔧 Intentando con LibreOffice...")
                    result = subprocess.run([
                        'libreoffice', '--headless', '--convert-to', 'docx', 
                        '--outdir', temp_dir, temp_input_path
                    ], capture_output=True, text=True, timeout=60)
                    
                    if result.returncode == 0 and os.path.exists(archivo_salida):
                        archivos_convertidos.append(archivo_salida)
                        print(f"✅ Convertido con LibreOffice: {archivo_salida}")
                        continue
                    else:
                        print("❌ LibreOffice no disponible o falló")
                except Exception as e3:
                    print(f"❌ Error con LibreOffice: {e3}")
            
            # Si se extrajo texto, crear documento Word
            if texto_extraido.strip():
                try:
                    from docx import Document
                    doc = Document()
                    
                    # Agregar título
                    doc.add_heading(f'Documento convertido: {file.filename}', 0)
                    doc.add_paragraph(f'Archivo original: {file.filename}')
                    doc.add_paragraph(f'Fecha de conversión: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}')
                    doc.add_paragraph(f'Método de conversión: {metodo_usado}')
                    doc.add_paragraph(f'Caracteres extraídos: {len(texto_extraido)}')
                    doc.add_paragraph('─' * 50)
                    
                    # Dividir el texto en párrafos y agregarlos al documento
                    parrafos = texto_extraido.split('\n\n')
                    for parrafo in parrafos:
                        if parrafo.strip():
                            # Limpiar el texto
                            parrafo_limpio = parrafo.strip().replace('\n', ' ')
                            if len(parrafo_limpio) > 0:
                                doc.add_paragraph(parrafo_limpio)
                    
                    # Guardar el documento
                    doc.save(archivo_salida)
                    archivos_convertidos.append(archivo_salida)
                    print(f"✅ Documento Word creado exitosamente: {archivo_salida}")
                    
                except ImportError:
                    # Si no está disponible python-docx, crear archivo de texto
                    archivo_salida_txt = archivo_salida.replace('.docx', '.txt')
                    with open(archivo_salida_txt, 'w', encoding='utf-8') as txt_file:
                        txt_file.write(f"DOCUMENTO CONVERTIDO: {file.filename}\n")
                        txt_file.write(f"MÉTODO: {metodo_usado}\n")
                        txt_file.write("=" * 50 + "\n\n")
                        txt_file.write(texto_extraido)
                    archivos_convertidos.append(archivo_salida_txt)
                    print(f"✅ Archivo de texto creado: {archivo_salida_txt}")
                    
            else:
                # Si no se pudo extraer texto, crear un documento informativo
                try:
                    from docx import Document
                    doc = Document()
                    doc.add_heading(f'Archivo PDF: {file.filename}', 0)
                    doc.add_paragraph('No se pudo extraer texto de este archivo PDF.')
                    doc.add_paragraph('Posibles causas:')
                    doc.add_paragraph('• El PDF está protegido con contraseña')
                    doc.add_paragraph('• El PDF contiene solo imágenes (escaneado) - requiere OCR')
                    doc.add_paragraph('• El archivo está corrupto')
                    doc.add_paragraph('• Formato no compatible')
                    doc.add_paragraph('• Dependencias OCR no instaladas')
                    doc.save(archivo_salida)
                    archivos_convertidos.append(archivo_salida)
                    print(f"⚠️ Documento informativo creado: {archivo_salida}")
                    
                except ImportError:
                    archivo_salida_txt = archivo_salida.replace('.docx', '.txt')
                    with open(archivo_salida_txt, 'w', encoding='utf-8') as txt_file:
                        txt_file.write(f"ARCHIVO PDF: {file.filename}\n")
                        txt_file.write("No se pudo extraer texto de este archivo PDF.\n")
                    archivos_convertidos.append(archivo_salida_txt)
                    
        except Exception as e:
            print(f"❌ Error general procesando {file.filename}: {str(e)}")
            continue
    
    if not archivos_convertidos:
        return JSONResponse(
            status_code=400,
            content={"error": "No se pudo convertir ningún archivo PDF. Verifique que los archivos sean PDFs válidos y no estén protegidos."}
        )
    
    # Crear ZIP con los archivos convertidos
    zip_path = os.path.join(temp_dir, "pdf_a_word_convertidos.zip")
    with zipfile.ZipFile(zip_path, "w") as zipf:
        for archivo in archivos_convertidos:
            zipf.write(archivo, os.path.basename(archivo))
    
    return FileResponse(
        zip_path,
        filename="pdf_a_word_convertidos.zip",
        media_type="application/zip"
    )
