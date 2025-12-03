#!/usr/bin/env python3
"""
Script para probar la implementación de OCR en PDFs
"""

import sys
import os
sys.path.append('/home/yeralcode/Documentos/ITRC/PROJECT/ExcelSior/backend')

import tempfile
import pdf2image
import pytesseract
from docx import Document

def extract_text_with_ocr(pdf_path):
    """Extrae texto de un PDF usando OCR"""
    print(f"🔍 Iniciando OCR para: {pdf_path}")
    
    texto_extraido = ""
    
    try:
        # Convertir PDF a imágenes
        images = pdf2image.convert_from_path(pdf_path)
        print(f"📄 PDF convertido a {len(images)} imágenes")
        
        for i, image in enumerate(images):
            print(f"   🔍 Procesando página {i+1} con OCR...")
            
            # Extraer texto usando Tesseract OCR
            page_text = pytesseract.image_to_string(image, lang='spa')
            
            if page_text.strip():
                texto_extraido += f"--- PÁGINA {i+1} ---\n{page_text}\n\n"
                print(f"   ✅ Página {i+1}: {len(page_text)} caracteres extraídos")
            else:
                print(f"   ⚠️  Página {i+1}: No se pudo extraer texto")
                
    except Exception as e:
        print(f"❌ Error en OCR: {e}")
        return ""
    
    return texto_extraido

def create_word_document(texto, output_path):
    """Crea un documento Word con el texto extraído"""
    try:
        doc = Document()
        doc.add_heading('Documento convertido con OCR', 0)
        doc.add_paragraph('Texto extraído de PDF escaneado usando Tesseract OCR')
        doc.add_paragraph('─' * 50)
        
        # Dividir el texto en párrafos
        lineas = texto.split('\n')
        for linea in lineas:
            if linea.strip() and not linea.startswith('--- PÁGINA'):
                doc.add_paragraph(linea.strip())
        
        doc.save(output_path)
        print(f"✅ Documento Word creado: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ Error creando documento Word: {e}")
        return False

if __name__ == "__main__":
    # Ruta del archivo PDF a procesar
    pdf_path = "/home/yeralcode/Descargas/Entrega de Compromisos ANI 171 Coljuegos Vulnerabilidades y riesgos en la operación de juegos de suerte y azar en línea_ gestión de bloqueos y control de operadores ilegales/SITIOS Y REDES COMUNICADOS DE COLJUEGOS A POLICIA/20245200208741.pdf"
    
    print("🧪 Probando implementación de OCR para PDFs escaneados")
    print("=" * 70)
    
    # Extraer texto con OCR
    texto_ocr = extract_text_with_ocr(pdf_path)
    
    if texto_ocr.strip():
        print(f"\n📊 Texto extraído con OCR: {len(texto_ocr)} caracteres")
        print("📝 Muestra del texto:")
        print("-" * 50)
        print(texto_ocr[:500] + "..." if len(texto_ocr) > 500 else texto_ocr)
        print("-" * 50)
        
        # Crear documento Word
        temp_dir = tempfile.mkdtemp()
        docx_path = os.path.join(temp_dir, "prueba_ocr.docx")
        
        if create_word_document(texto_ocr, docx_path):
            print(f"\n🎉 ¡OCR funcionó correctamente!")
            print(f"📁 Documento guardado en: {docx_path}")
        else:
            print("\n❌ Falló la creación del documento Word")
            
    else:
        print("\n❌ No se pudo extraer texto con OCR")
